import warnings
warnings.filterwarnings("ignore") 
import pandas as pd
from clean_data import text_prepare
import argparse
from transformers import AutoTokenizer, AutoModel, set_seed as set_seed_transformers
import torch
print("Cuda available: ", torch.cuda.is_available())
import torch.nn as nn
from torch.optim import Adam
import numpy as np
from sklearn.model_selection import train_test_split
from torch.utils.data import Dataset, DataLoader
from sklearn.metrics import classification_report
from sklearn.preprocessing import LabelEncoder
import os
import time
from replace_with_rbf import replace_ffn_with_rbf_bert
from get_dataset import download_prepare_dataset
from create_plots import create_plots
from transformers import logging
import requests
import random
import sys
from io import StringIO
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
output = StringIO()
sys.stdout = output

logging.set_verbosity_error()

def set_seed(seed):
    random.seed(seed)
    np.random.seed(seed)
    set_seed_transformers(seed)
    torch.manual_seed(seed)
    if torch.cuda.is_available():
        torch.cuda.manual_seed_all(seed)

set_seed(42)

def print_number_of_trainable_model_parameters(model):
    trainable_model_params = 0
    all_model_params = 0
    for _, param in model.named_parameters():
        all_model_params += param.numel()
        if param.requires_grad:
            trainable_model_params += param.numel()
    return f"trainable model parameters: {trainable_model_params}\nall model parameters: {all_model_params}\npercentage of trainable model parameters: {100 * trainable_model_params / all_model_params:.2f}%"

def get_model_tokenizer(model_id = "lanwuwei/GigaBERT-v3-Arabic-and-English"):
    tokenizer = AutoTokenizer.from_pretrained(model_id)
    model = AutoModel.from_pretrained(model_id)
    return model, tokenizer

class MyModel(nn.Module):
    def __init__(self, bert, num_classes = 2):

        super(MyModel, self).__init__()

        self.bert = bert
        self.dropout = nn.Dropout(0.25)
        self.linear1 = nn.Linear(768, 384)
        self.linear2 = nn.Linear(384, num_classes)
        # self.softmax = nn.Softmax()

    def forward(self, input_ids, attention_mask):
        pooled_output = self.bert(input_ids, attention_mask)[0][:,0]
        output = self.linear1(pooled_output)
        output = self.dropout(output)
        output = self.linear2(output)
        # output = self.softmax(output)
        return output
    
class dataset(Dataset):
    def __init__(self, X, Y, tokenizer, device):
        self.X = [tokenizer(x,
                            max_length = 150,
                            truncation = True,
                            padding = 'max_length',
                            return_tensors='pt').to(device)
                for x in X
                ]

        self.Y = torch.tensor(Y, dtype = torch.long).to(device)
    def __len__(self):
        return len(self.X)
    def __getitem__(self, index):
        return self.X[index], self.Y[index]
    
def train(args):
    le= LabelEncoder()
    device = args.device
    data_df = download_prepare_dataset()
    print("Data Distribution: ")
    print(data_df['category'].value_counts())
    data_df['category'] = le.fit_transform(data_df['category'])
    print()
    print("Classes are: ", len(le.classes_) , " which are: " ,  le.classes_)
    args.save_model_path = f"{f'RBF_{args.num_kernels}_{args.kernel_name}-kernels' if args.enable_rbf else 'Basic'}_Epochs_{args.EPOCHS}"
    if os.path.exists(args.save_model_path):
        print(f" {args.save_model_path} Directory already exists")
    else:
        os.makedirs(args.save_model_path)
        print(f"{args.save_model_path} Directory Created")
    np.save(os.path.join(args.save_model_path, "label_encoder.npy"), le.classes_)

    testing_predictions = []
    model, tokenizer = get_model_tokenizer(args.model_id)
    # print(print_number_of_trainable_model_parameters(model))
    

    for param in model.parameters():
        param.requires_grad = False
    model = MyModel(model, len(data_df['category'].unique())).to(device)
    print_number_of_trainable_model_parameters(model)
    print()
    if args.enable_rbf:
        print("Number of Kernels: ", args.num_kernels)
        print()
        replace_ffn_with_rbf_bert(model, num_kernels=args.num_kernels, kernel_name=args.kernel_name)
        print()
        print(f"Feedforward layers in BERT have been replaced with RBF layers of {args.num_kernels} {args.kernel_name} kernels")
        print()
        print("After RBF Replacement model size becomes\n", print_number_of_trainable_model_parameters(model))
    print("=="*30)
    
    criterion = nn.CrossEntropyLoss()
    optimizer = Adam(model.parameters(), lr= args.lr)

    X_train, X_test, y_train, y_test = train_test_split(np.array(data_df["text"]), np.array(data_df["category"]), test_size=0.3)
    X_val, X_test, y_val, y_test = train_test_split(X_test, y_test, test_size=0.5)
    print("Training set is: ", X_train.shape[0], " rows which is ", round(X_train.shape[0]/data_df.shape[0],4)*100, "%")
    print("Validation set is: ",X_val.shape[0], " rows which is ", round(X_val.shape[0]/data_df.shape[0],4)*100, "%")
    print("Testing set is: ",X_test.shape[0], " rows which is ", round(X_test.shape[0]/data_df.shape[0],4)*100, "%")

    training_data = dataset(X_train, y_train, tokenizer, device)
    validation_data = dataset(X_val, y_val, tokenizer, device)
    testing_data = dataset(X_test, y_test, tokenizer, device)

    train_dataloader = DataLoader(training_data, batch_size=args.BATCH_SIZE, shuffle= False)
    validation_dataloader = DataLoader(validation_data, batch_size=args.BATCH_SIZE, shuffle= False)
    testing_dataloader = DataLoader(testing_data, batch_size=args.BATCH_SIZE, shuffle= False)

    total_loss_train_plot = []
    total_loss_validation_plot = []
    total_acc_train_plot = []
    total_acc_validation_plot = []
    t1 = time.time()
    for epoch in range(args.EPOCHS):
        total_acc_train = 0
        total_loss_train = 0
        total_acc_val = 0
        total_loss_val = 0
        ## Training and Validation
        for indx, data in enumerate(train_dataloader):
            input, label = data

            input.to(device)
            label.to(device)

            prediction = model(input['input_ids'].squeeze(1),
                                input['attention_mask'].squeeze(1)).squeeze(1)

            batch_loss = criterion(prediction, label)

            total_loss_train += batch_loss.item()

            acc = (prediction.argmax(dim=1).round() == label).sum().item()
            total_acc_train += acc

            batch_loss.backward()
            optimizer.step()
            optimizer.zero_grad()

        ## Validation
        with torch.no_grad():
            for indx, data in enumerate(validation_dataloader):
                input, label = data
                input.to(device)
                label.to(device)

                prediction = model(input['input_ids'].squeeze(1),
                                input['attention_mask'].squeeze(1)).squeeze(1)

                batch_loss_val = criterion(prediction, label)
                total_loss_val += batch_loss_val.item()


                acc = (prediction.argmax(dim=1).round() == label).sum().item()

                total_acc_val += acc


        total_loss_train_plot.append(round(total_loss_train/1000, 4))
        total_loss_validation_plot.append(round(total_loss_val/100, 4))
        total_acc_train_plot.append(round(total_acc_train/(training_data.__len__())*100, 4))
        total_acc_validation_plot.append(round(total_acc_val/(validation_data.__len__())*100, 4))

        print(f'''Epoch no. {epoch + 1} Train Loss: {total_loss_train/1000:.4f} Train Accuracy: {(total_acc_train/(training_data.__len__())*100):.4f} Validation Loss: {total_loss_val/100:.4f} Validation Accuracy: {(total_acc_val/(validation_data.__len__())*100):.4f}''')
        print("="*50)
    print(f"Training Time: {time.time()-t1:.2f} seconds")
    create_plots(total_loss_train_plot, 
                    total_loss_validation_plot, 
                    total_acc_train_plot, 
                    total_acc_validation_plot, 
                    args.EPOCHS, 
                    os.path.join(args.save_model_path, f"{f'RBF_{args.num_kernels}_{args.kernel_name}-kernels' if args.enable_rbf else 'Basic'}_Epochs_{args.EPOCHS}_plot.png"))
    # Model Save
    torch.save(model.state_dict(), os.path.join(args.save_model_path, "model.pth"))
    print("Model Saved")
    t1 = time.time()
    with torch.no_grad():
        total_loss_test = 0
        total_acc_test = 0
        for indx, data in enumerate(testing_dataloader):
            input, label = data
            input.to(device)
            label.to(device)

            prediction = model(input['input_ids'].squeeze(1), input['attention_mask'].squeeze(1)).squeeze(1)
            batch_loss_val = criterion(prediction, label)
            total_loss_test += batch_loss_val.item()
            acc = (prediction.argmax(dim=1).round() == label).sum().item()
            testing_predictions += prediction.argmax(dim=1).round().cpu().numpy().tolist()
            total_acc_test += acc
    print(f"Testing Time: {time.time()-t1:.2f} seconds")

    print(f"Accuracy Score is: {round((total_acc_test/X_test.shape[0])*100, 2)}%")
    print("Classification Report: ")
    print(classification_report(testing_data.Y.cpu().numpy() , testing_predictions, target_names=le.classes_))
    requests.post("https://ntfy.sh/rbf_bert", data=f"{args.save_model_path} is Done".encode(encoding='utf-8'))

    sys.stdout = sys.__stdout__
    printed_text = output.getvalue()
    doc = Document()
    doc.add_paragraph(printed_text)
    # Insert the plot at the end of the document
    doc.add_picture(os.path.join(args.save_model_path, f"{f'RBF_{args.num_kernels}_{args.kernel_name}-kernels' if args.enable_rbf else 'Basic'}_Epochs_{args.EPOCHS}_plot.png"), width=Inches(5))
    doc.save(f"{args.save_model_path}.docx")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--model_id", type=str, default="lanwuwei/GigaBERT-v3-Arabic-and-English")
    parser.add_argument("--device", type=str, default="cuda" if torch.cuda.is_available() else "cpu")
    parser.add_argument("--lr", type=float, default=1e-4)
    parser.add_argument("--BATCH_SIZE", type=int, default=16)
    parser.add_argument("--EPOCHS", type=int, default=10)
    parser.add_argument("--save_model_path", type=str, default="basic_model_path")
    parser.add_argument("--num_kernels", type=int, default=2)
    parser.add_argument("--kernel_name", type=str, default="linear")
    parser.add_argument("--enable_rbf", type=bool, default=False)
    args = parser.parse_args()
    train(args)